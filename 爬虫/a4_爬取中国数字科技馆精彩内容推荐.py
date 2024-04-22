# 请求地址
import re

import requests
from bs4 import BeautifulSoup
import docx

url = "https://www.cdstm.cn/subjects/jcnrtj/202204/t20220406_1067359.html"
headers = {
    'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36 Edg/115.0.1901.188'}
response = requests.get(url, headers=headers)
response.encoding = response.apparent_encoding
soup = BeautifulSoup(response.text, "html.parser")
article = soup.find("div", class_="article")
# print(article)
# 处理数据
doc = docx.Document()
# 大标题  时间  文章主题
article_headline = article.find("h1", class_='article-title').text
# print(article_headline)
article_time = article.find("div", class_='time').find('time').text
# print(article_time)
article_content = article.find('div', class_='article-content').find('div', class_='gxchange').find('div',
                                                                                                    class_='TRS_Editor')
print(article_content)
paragraphs = article_content.find_all("p")
# 写出数据
doc.add_heading(article_headline, level=0)
for paragraph in paragraphs:
    if paragraph.find('b') is not None:
        doc.add_heading(paragraph.text, level=1)
    else:
        if paragraph.find("img") is not None:
            pic_url = paragraph.find('img')['src']
            print(pic_url)
            continue
        if paragraph.find('font') is not None:
            continue
        doc.add_paragraph(paragraph.text)
doc.save("my_doc.docx")

